"""
Document extraction module for processing DOCX files with comprehensive error handling and logging.
"""
import json
import os
import traceback
from typing import Dict, List, Any, Optional
import docx

from utils.logger import get_logger

# Initialize logger
logger = get_logger(__name__)

class DocxTopicExtractor:
    """Extracts topics and subtopics from DOCX files with structured content."""
    
    def __init__(self):
        """Initialize the extractor with an empty content dictionary."""
        self.extracted_content = {}
    
    def extract_topics_and_subtopics(self, file_path: str, topic_filter: List[str] = None) -> Dict[str, Any]:
        """
        Extract topics (Heading 2) and subtopics (Heading 3) from DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            topic_filter: Optional list of specific topic titles to extract
        
        Returns:
            Dictionary with topics and their subtopics with content
            
        Raises:
            FileNotFoundError: If the input file doesn't exist
            docx.opc.exceptions.PackageNotFoundError: If the file is not a valid DOCX
            Exception: For other unexpected errors
        """
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
            
        logger.info(f"Starting extraction from {file_path}")
        
        try:
            doc = docx.Document(file_path)
            structure = {}
            current_topic = None
            current_subtopic = None
            current_text = []
            
            # Convert filter to lowercase for case-insensitive matching
            topic_filter_lower = [title.lower() for title in topic_filter] if topic_filter else None
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                try:
                    paragraph_text = paragraph.text.strip()
                    
                    # Skip empty paragraphs
                    if not paragraph_text:
                        continue
                    
                    # Main topic (Heading 2)
                    if paragraph.style.name == 'Heading 2':
                        # Save previous content before switching topics
                        if current_topic:
                            self._save_content(structure, current_topic, current_subtopic, current_text)
                        
                        # Check if this topic should be extracted
                        if topic_filter_lower and paragraph_text.lower() not in topic_filter_lower:
                            current_topic = None
                            current_subtopic = None
                            current_text = []
                            logger.debug(f"Skipping filtered out topic: {paragraph_text}")
                            continue
                        
                        # Start new topic
                        current_topic = paragraph_text
                        current_subtopic = None
                        current_text = []
                        structure[current_topic] = {'text': '', 'subtopics': {}}
                        logger.debug(f"New topic found: {current_topic}")
                    
                    # Subtopic (Heading 3)
                    elif paragraph.style.name == 'Heading 3' and current_topic:
                        # Save previous subtopic content
                        if current_subtopic:
                            structure[current_topic]['subtopics'][current_subtopic] = {'text': '\n'.join(current_text).strip()}
                        
                        # Start new subtopic
                        current_subtopic = paragraph_text
                        current_text = []
                        logger.debug(f"New subtopic found under '{current_topic}': {current_subtopic}")
                    
                    # Regular content - only collect if we're tracking a topic
                    elif current_topic and paragraph.style.name not in ['TOC 1', 'TOC 2', 'TOC 3', 'Header', 'Footer']:
                        current_text.append(paragraph_text)
                        
                except Exception as e:
                    logger.error(f"Error processing paragraph {para_num}: {str(e)}")
                    logger.debug(traceback.format_exc())
                    continue
            
            # Save the last content
            if current_topic:
                self._save_content(structure, current_topic, current_subtopic, current_text)
            
            logger.info(f"Successfully extracted content from {file_path}")
            return structure
            
        except docx.opc.exceptions.PackageNotFoundError:
            error_msg = f"Invalid or corrupted DOCX file: {file_path}"
            logger.error(error_msg)
            raise
        except Exception as e:
            error_msg = f"Unexpected error processing {file_path}: {str(e)}"
            logger.error(error_msg)
            logger.debug(traceback.format_exc())
            raise
    
    def _save_content(self, structure: Dict, current_topic: str, current_subtopic: Optional[str], current_text: List[str]) -> None:
        """
        Save current content to the appropriate place in structure.
        
        Args:
            structure: The main structure dictionary
            current_topic: Current topic name
            current_subtopic: Current subtopic name (can be None)
            current_text: List of text lines to save
        """
        try:
            content = '\n'.join(current_text).strip()
            
            if current_subtopic:
                # Save to subtopic with 'text' key
                structure[current_topic]['subtopics'][current_subtopic] = {'text': content}
                logger.debug(f"Saved content to subtopic: {current_topic} > {current_subtopic}")
            else:
                # Save to main topic text
                structure[current_topic]['text'] = content
                logger.debug(f"Saved content to topic: {current_topic}")
                
        except Exception as e:
            logger.error(f"Error saving content: {str(e)}")
            logger.debug(traceback.format_exc())
    
    def analyze_document_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Analyze the document to show heading structure.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with analysis of the document structure
        """
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
            
        logger.info(f"Analyzing document structure: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            analysis = {
                'heading_counts': {},
                'topics': [],
                'structure_preview': []
            }
            
            current_topic = None
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                try:
                    if paragraph.style.name.startswith('Heading') and paragraph.text.strip():
                        style = paragraph.style.name
                        text = paragraph.text.strip()
                        
                        # Count heading types
                        if style not in analysis['heading_counts']:
                            analysis['heading_counts'][style] = 0
                        analysis['heading_counts'][style] += 1
                        
                        # Track structure
                        if style == 'Heading 2':
                            current_topic = text
                            analysis['topics'].append(text)
                            analysis['structure_preview'].append(f" {text}")
                            logger.debug(f"Found topic: {text}")
                        elif style == 'Heading 3' and current_topic:
                            analysis['structure_preview'].append(f"   {text}")
                            logger.debug(f"Found subtopic under {current_topic}: {text}")
                except Exception as e:
                    logger.error(f"Error analyzing paragraph {para_num}: {str(e)}")
                    logger.debug(traceback.format_exc())
            
            logger.info(f"Document analysis complete. Found {len(analysis['topics'])} topics.")
            return analysis
            
        except Exception as e:
            error_msg = f"Error analyzing document: {str(e)}"
            logger.error(error_msg)
            logger.debug(traceback.format_exc())
            raise
    
    def save_to_json(self, structure: Dict[str, Any], output_file: str) -> None:
        """
        Save extracted content to JSON file.
        
        Args:
            structure: The data structure to save
            output_file: Path to the output JSON file
            
        Raises:
            IOError: If there's an error writing the file
        """
        try:
            os.makedirs(os.path.dirname(os.path.abspath(output_file)), exist_ok=True)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(structure, f, indent=2, ensure_ascii=False)
                
            logger.info(f"Content successfully saved to '{output_file}'")
            
        except Exception as e:
            error_msg = f"Error saving to JSON file {output_file}: {str(e)}"
            logger.error(error_msg)
            logger.debug(traceback.format_exc())
            raise IOError(error_msg)
    
    def print_extraction_summary(self, structure: Dict[str, Any]) -> None:
        """Print a summary of extracted content."""
        if not structure:
            logger.warning("No content to summarize (empty structure)")
            return
            
        try:
            print("\n" + "="*60)
            print(" EXTRACTION SUMMARY")
            print("="*60)
            
            topic_count = len(structure)
            subtopic_count = sum(len(content['subtopics']) for content in structure.values())
            
            print(f" Total Topics: {topic_count}")
            print(f" Total Subtopics: {subtopic_count}")
            
            print("\n Content Details:")
            for topic_name, content in structure.items():
                main_text_chars = len(content['text'])
                subtopic_count = len(content['subtopics'])
                
                print(f"\n '{topic_name}'")
                print(f"   Main content: {main_text_chars} characters")
                print(f"   Subtopics: {subtopic_count}")
                
                for subtopic_name, subtopic_text in content['subtopics'].items():
                    print(f"      {subtopic_name}: {len(subtopic_text['text'])} characters")
                    
            logger.info(f"Printed extraction summary: {topic_count} topics, {subtopic_count} subtopics")
            
        except Exception as e:
            logger.error(f"Error generating extraction summary: {str(e)}")
            logger.debug(traceback.format_exc())
            print("Error generating extraction summary. Check logs for details.")
    
    def print_content_preview(self, structure: Dict[str, Any], preview_length: int = 150) -> None:
        """Print content with preview text."""
        if not structure:
            logger.warning("No content to preview (empty structure)")
            return
            
        try:
            print("\n" + "="*60)
            print(" CONTENT PREVIEW")
            print("="*60)
            
            for topic_name, content in structure.items():
                print(f"\n TOPIC: {topic_name}")
                print("-" * 50)
                
                if content['text']:
                    preview = content['text'][:preview_length]
                    if len(content['text']) > preview_length:
                        preview += "..."
                    print(f"Main Text: {preview}")
                
                if content['subtopics']:
                    print(f"\nSubtopics:")
                    for subtopic_name, subtopic_text in content['subtopics'].items():
                        preview = subtopic_text['text'][:preview_length]
                        if len(subtopic_text['text']) > preview_length:
                            preview += "..."
                        print(f"  {subtopic_name}")
                        print(f"     {preview}")
                        
            logger.info("Printed content preview")
            
        except Exception as e:
            logger.error(f"Error generating content preview: {str(e)}")
            logger.debug(traceback.format_exc())
            print("Error generating content preview. Check logs for details.")


def extract_json_from_doc(path: str, output_file: Optional[str] = None) -> Dict[str, Any]:
    """
    Extract JSON structure from a DOCX file and optionally save to a file.
    
    Args:
        path: Path to the DOCX file
        output_file: Optional path to save the extracted JSON
        
    Returns:
        Extracted content as a dictionary
    """
    try:
        extractor = DocxTopicExtractor()
        structure = extractor.extract_topics_and_subtopics(path)
        
        # Save to file if output path is provided
        if output_file:
            extractor.save_to_json(structure, output_file)
            logger.info(f"Extracted data saved to {os.path.abspath(output_file)}")
            
        return structure
    except Exception as e:
        logger.error(f"Error in extract_json_from_doc: {str(e)}")
        logger.debug(traceback.format_exc())
        raise